from __future__ import annotations

import hashlib
import json
import re
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Any

from testbox.sdk import PluginError, Result

ROLES = {
    "测试名称": ("测试名称", "用例名", "用例名称", "测试用例", "case name", "test name", "name"),
    "验证点": ("验证点", "检查点", "校验点", "检查项", "checkpoint", "check point"),
    "步骤名称": ("步骤名称", "步骤", "操作步骤", "step name", "step", "步骤标题"),
    "步骤描述": ("步骤描述", "步骤说明", "操作说明", "描述", "step description", "desc"),
    "预期结果": ("预期结果", "期望结果", "expected result", "expected", "期望"),
    "测试结果": ("测试结果", "结果", "状态", "执行结果", "result", "status", "pass/fail"),
}
REQUIRED = ("测试名称", "验证点", "测试结果")
STEP_ROLES = ("步骤名称", "步骤描述", "预期结果")
EXECUTED = {"已执行", "pass", "passed", "通过", "done", "completed", "完成", "成功"}
NORMALIZE_RE = re.compile(r"[\s\-_·./\\|:：,，。；;()（）\[\]{}<>《》\"'`~!@#$%^&*=+?]+")
INVALID_FILENAME_RE = re.compile(r'[\\/:*?"<>|\r\n\t]+')


def normalize(value: object) -> str:
    return NORMALIZE_RE.sub("", str(value or "").strip().lower())


def suggest_mapping(headers: list[str]) -> tuple[dict[str, str | None], dict[str, str]]:
    normalized = {normalize(header): header for header in headers}
    mapping: dict[str, str | None] = {}
    reasons: dict[str, str] = {}
    for role, aliases in ROLES.items():
        match = next((normalized[normalize(alias)] for alias in aliases if normalize(alias) in normalized), None)
        reason = f"自动识别为：{match}" if match else "未识别，请手动选择"
        if match is None:
            for header in headers:
                probe = normalize(header)
                if probe and any(normalize(alias) in probe or probe in normalize(alias) for alias in aliases):
                    match = header
                    reason = f"智能匹配为：{match}"
                    break
        mapping[role] = match
        reasons[role] = reason
    return mapping, reasons


@dataclass(frozen=True)
class ColumnInfo:
    index: int
    header: str
    samples: tuple[str, ...]


@dataclass(frozen=True)
class TestItem:
    row_index: int
    case_name: str
    checkpoint: str
    step_name: str = ""
    step_description: str = ""
    expected_result: str = ""

    @property
    def is_step(self) -> bool:
        return bool(self.step_name or self.step_description or self.expected_result)

    @property
    def display_title(self) -> str:
        return f"{self.case_name} / {self.step_name}" if self.is_step and self.step_name else self.case_name

    def step_note(self) -> str:
        pairs = (("步骤名称", self.step_name), ("步骤描述", self.step_description), ("预期结果", self.expected_result))
        return "\n".join(f"{label}：{value}" for label, value in pairs if value)


class WorkbookCases:
    """The original screenshot-to-word Excel workflow, adapted to task outputs."""

    def __init__(self, path: Path, mapping: dict[str, str | None] | None, scan_rows: int, *, require_mapping: bool = True):
        try:
            from openpyxl import load_workbook
        except ModuleNotFoundError as error:
            raise PluginError("DEPENDENCY_MISSING", "Evidence Tool 需要 openpyxl，请安装插件 requirements.txt") from error
        if path.suffix.lower() not in {".xlsx", ".xlsm"}:
            raise PluginError("INPUT_INVALID", "Excel 用例文件只支持 .xlsx 或 .xlsm")
        try:
            self.workbook = load_workbook(path, keep_vba=path.suffix.lower() == ".xlsm")
        except Exception as error:
            raise PluginError("INPUT_INVALID", "无法读取 Excel，仅支持有效的 .xlsx 或 .xlsm 文件") from error
        self.path = path
        self.sheet = self.workbook.active
        self.header_row = self._detect_header(scan_rows)
        self.columns_info = self._columns()
        self.headers = [column.header for column in self.columns_info]
        self.suggested_mapping, self.mapping_reasons = suggest_mapping(self.headers)
        selected = mapping if mapping is not None else self.suggested_mapping
        positions = {header: index + 1 for index, header in enumerate(self.headers)}
        self.columns = {role: positions[header] for role, header in selected.items() if header in positions}
        self.mapping = {role: self.headers[column - 1] for role, column in self.columns.items()}
        self._merged: dict[tuple[int, int], tuple[int, int]] = {}
        if require_mapping:
            self.ensure_required_mapping()

    def ensure_required_mapping(self) -> None:
        missing = [role for role in REQUIRED if role not in self.columns]
        if missing:
            raise PluginError(
                "COLUMN_MAPPING_REQUIRED",
                f"Excel 缺少必要列映射：{'、'.join(missing)}",
                details=self.preview_data(missing),
            )

    def preview_data(self, missing: list[str] | None = None) -> dict[str, Any]:
        return {
            "sheet": self.sheet.title,
            "header_row": self.header_row,
            "columns": [asdict(column) for column in self.columns_info],
            "mapping": self.mapping,
            "suggested_mapping": self.suggested_mapping,
            "mapping_reasons": self.mapping_reasons,
            "required_roles": list(REQUIRED),
            "missing_required_roles": missing if missing is not None else [role for role in REQUIRED if role not in self.columns],
        }

    def _detect_header(self, scan_rows: int) -> int:
        aliases = {normalize(alias) for values in ROLES.values() for alias in values}
        best_score, best_row = -1, 1
        for row in range(1, min(self.sheet.max_row, scan_rows) + 1):
            values = [self.sheet.cell(row, column).value for column in range(1, self.sheet.max_column + 1)]
            non_empty = [str(value).strip() for value in values if value not in (None, "")]
            if len(non_empty) < 2:
                continue
            score = len(non_empty) + len({normalize(value) for value in non_empty} & aliases) * 5
            if score > best_score:
                best_score, best_row = score, row
        return best_row

    def _columns(self) -> list[ColumnInfo]:
        columns: list[ColumnInfo] = []
        used: set[str] = set()
        for column in range(1, self.sheet.max_column + 1):
            value = self.sheet.cell(self.header_row, column).value
            header = str(value).strip() if value not in (None, "") else f"第{column}列"
            if header in used:
                header = f"{header}({column})"
            used.add(header)
            samples: list[str] = []
            for row in range(self.header_row + 1, min(self.sheet.max_row, self.header_row + 5) + 1):
                value = self.sheet.cell(row, column).value
                if value not in (None, ""):
                    samples.append(str(value).strip())
            columns.append(ColumnInfo(column, header, tuple(samples[:3])))
        return columns

    def _parent(self, row: int, column: int) -> tuple[int, int]:
        key = (row, column)
        if key not in self._merged:
            self._merged[key] = next(
                ((area.min_row, area.min_col) for area in self.sheet.merged_cells.ranges if area.min_row <= row <= area.max_row and area.min_col <= column <= area.max_col),
                key,
            )
        return self._merged[key]

    def text(self, row: int, role: str, *, resolve_merged: bool = False) -> str:
        if role not in self.columns:
            return ""
        column = self.columns[role]
        value = self.sheet.cell(row, column).value
        if value in (None, "") and resolve_merged:
            value = self.sheet.cell(*self._parent(row, column)).value
        return str(value).strip() if value not in (None, "") else ""

    def pending(self) -> list[TestItem]:
        items: list[TestItem] = []
        current_case = ""
        current_checkpoint = ""
        executed = {normalize(value) for value in EXECUTED}
        for row in range(self.header_row + 1, self.sheet.max_row + 1):
            raw_case = self.text(row, "测试名称", resolve_merged=True)
            raw_checkpoint = self.text(row, "验证点", resolve_merged=True)
            if raw_case and raw_case != current_case:
                # This mirrors the source project: a new case cannot inherit a
                # checkpoint from the previous case when its own cell is blank.
                current_case = raw_case
                current_checkpoint = raw_checkpoint or ""
            elif raw_case:
                current_case = raw_case
            if raw_checkpoint:
                current_checkpoint = raw_checkpoint
            step = tuple(self.text(row, role) for role in STEP_ROLES)
            if not any((current_case, current_checkpoint, *step)):
                continue
            if normalize(self.text(row, "测试结果")) in executed:
                continue
            items.append(TestItem(row, current_case or f"未命名用例_{row}", current_checkpoint or "未填写验证点", *step))
        return items

    def mark(self, row: int, status: str) -> None:
        column = self.columns["测试结果"]
        self.sheet.cell(*self._parent(row, column)).value = status

    def save(self, path: Path) -> None:
        path.parent.mkdir(parents=True, exist_ok=True)
        self.workbook.save(path)

    def close(self) -> None:
        self.workbook.close()


def safe_filename(name: str) -> str:
    return INVALID_FILENAME_RE.sub("_", name).strip(" ._")[:80] or "未命名用例"


def report_names(case_names: list[str]) -> dict[str, str]:
    return {case_name: f"{safe_filename(case_name)[:71]}-{hashlib.sha256(case_name.encode('utf-8')).hexdigest()[:8]}.docx" for case_name in case_names}


def validate_screenshot(path: Path) -> None:
    if not path.is_file():
        raise PluginError("INPUT_INVALID", f"截图文件不存在：{path.name}")
    try:
        from PIL import Image
        with Image.open(path) as image:
            image.verify()
    except ModuleNotFoundError as error:
        raise PluginError("DEPENDENCY_MISSING", "Evidence Tool 需要 Pillow，请安装插件 requirements.txt") from error
    except Exception as error:
        raise PluginError("INPUT_INVALID", f"无法读取截图：{path.name}") from error


def build_word(path: Path, case_name: str, item: TestItem, image: Path, width: float, existing: Path | None = None) -> None:
    try:
        from docx import Document
        from docx.shared import Inches
    except ModuleNotFoundError as error:
        raise PluginError("DEPENDENCY_MISSING", "Evidence Tool 需要 python-docx，请安装插件 requirements.txt") from error
    try:
        doc = Document(str(existing)) if existing else Document()
    except Exception as error:
        raise PluginError("REPORT_APPEND_FAILED", f"无法打开已有 Word 报告：{existing.name if existing else path.name}") from error
    if not existing:
        doc.add_heading(case_name, level=1)
    checkpoints = {paragraph.text.removeprefix("验证点：").strip() for paragraph in doc.paragraphs if paragraph.text.strip().startswith("验证点：")}
    if item.checkpoint not in checkpoints:
        doc.add_paragraph(f"验证点：{item.checkpoint}")
    for line in item.step_note().splitlines():
        if line.strip():
            doc.add_paragraph(line.strip())
    doc.add_picture(str(image), width=Inches(width))
    doc.add_paragraph()
    temporary = path.with_suffix(path.suffix + ".tmp")
    try:
        doc.save(temporary)
        temporary.replace(path)
    except Exception as error:
        raise PluginError("WORD_WRITE_FAILED", f"无法写入 Word 报告：{path.name}") from error


class Plugin:
    def init(self, context):
        self.context = context

    def _load_interactive_module(self):
        import importlib.util

        source = Path(__file__).with_name("interactive.py")
        spec = importlib.util.spec_from_file_location("evidence_tool_interactive", source)
        if spec is None or spec.loader is None:
            raise PluginError("INTERACTIVE_UNAVAILABLE", "无法加载交互截图模块")
        module = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(module)
        return module

    def execute(self, command, params):
        if command != "evidence.build":
            raise PluginError("INVALID_PARAMS", "不支持的命令")
        interactive = bool(params.get("interactive", True))
        if interactive and params.get("screenshots"):
            raise PluginError("INVALID_PARAMS", "交互截图模式不接收 screenshots；请删除截图文件，或将 interactive 设为 false 使用批处理模式")
        if not interactive and "screenshots" not in params:
            raise PluginError("INVALID_PARAMS", "批处理模式必须提供 screenshots")

        scan_rows = int(self.context.config.get("header_scan_rows", 15))
        requested_mapping = params.get("column_mapping")
        preview: WorkbookCases | None = None
        interactive_module = None
        try:
            if interactive and requested_mapping is None:
                # The source app always lets the user confirm its Excel mapping
                # before beginning execution. Keep this inside the plugin Host so
                # TestBox Core/GUI do not need a second evidence-specific flow.
                preview = WorkbookCases(Path(params["input"]), None, scan_rows, require_mapping=False)
                interactive_module = self._load_interactive_module()
                requested_mapping = interactive_module.confirm_mapping(preview.preview_data(), self.context.logger)
                if requested_mapping is None:
                    return Result("cancelled", "已取消 Excel 列映射", preview.preview_data(), [], [])
            cases = WorkbookCases(Path(params["input"]), requested_mapping, scan_rows)
        finally:
            if preview is not None:
                preview.close()

        try:
            items = cases.pending()
            if not items:
                summary = {**cases.preview_data(), "mode": "步骤版" if any(role in cases.columns for role in STEP_ROLES) else "基础版", "evidence_count": 0, "pending_count": 0, "items": [], "unmatched": []}
                self.context.files.write_text("evidence-index.json", json.dumps(summary, ensure_ascii=False, indent=2))
                return Result("success", "所有用例均已执行，或 Excel 中没有待执行数据", summary, ["evidence-index.json"], [])

            mode = "步骤版" if any(role in cases.columns for role in STEP_ROLES) else "基础版"
            width = float(params.get("image_width_inches", self.context.config.get("image_width_inches", 5.8)))
            if not 1.0 <= width <= 10.0:
                raise PluginError("INVALID_PARAMS", "image_width_inches 必须在 1 到 10 英寸之间")
            names = report_names([item.case_name for item in items])
            existing_reports = {Path(path).name: Path(path) for path in params.get("existing_reports", [])}
            status = params.get("status") or self.context.config.get("executed_status", "已执行")
            update_excel = bool(params.get("update_excel", True))
            excel_relative = f"executed-{Path(params['input']).name}"
            excel_target = self.context.files.resolve(excel_relative)
            saved: list[tuple[TestItem, Path]] = []
            reports: list[str] = []
            report_set: set[str] = set()

            def commit(item: TestItem, screenshot: Path) -> bool:
                validate_screenshot(screenshot)
                report_relative = f"reports/{names[item.case_name]}"
                report_target = self.context.files.resolve(report_relative)
                report_target.parent.mkdir(parents=True, exist_ok=True)
                source_report = report_target if report_target.exists() else existing_reports.get(report_target.name)
                build_word(report_target, item.case_name, item, screenshot, width, source_report)
                if report_relative not in report_set:
                    report_set.add(report_relative)
                    reports.append(report_relative)
                if update_excel:
                    cases.mark(item.row_index, status)
                    cases.save(excel_target)
                saved.append((item, screenshot))
                self.context.logger.info(f"已写入证据：第 {item.row_index} 行 / {item.display_title}")
                return True

            skipped: list[TestItem] = []
            ended = False
            screenshot_files: list[str] = []
            if interactive:
                module = interactive_module or self._load_interactive_module()
                screenshot_dir = self.context.files.resolve("screenshots")
                screenshot_dir.mkdir(parents=True, exist_ok=True)

                def commit_interactive(item: TestItem, screenshot: Path) -> bool:
                    # Keep the staged capture private until annotation is saved, then
                    # promote it to a stable task artifact before committing the report.
                    target = screenshot_dir / f"row-{item.row_index}.png"
                    try:
                        screenshot.replace(target)
                    except OSError:
                        import shutil
                        shutil.copy2(screenshot, target)
                    return commit(item, target)

                session = module.run_session(items, screenshot_dir, self.context.logger, commit_interactive)
                skipped = session["skipped"]
                ended = bool(session["ended"])
                screenshot_files = [f"screenshots/{path.name}" for _, path in saved]
            else:
                screenshots = [Path(path) for path in params["screenshots"]]
                for screenshot in screenshots:
                    validate_screenshot(screenshot)
                requested_rows = params.get("row_indexes")
                if requested_rows is not None:
                    if len(requested_rows) != len(screenshots) or len(set(requested_rows)) != len(requested_rows):
                        raise PluginError("INVALID_PARAMS", "row_indexes 必须与 screenshots 一一对应，且不能重复")
                    by_row = {item.row_index: item for item in items}
                    missing_rows = [row for row in requested_rows if row not in by_row]
                    if missing_rows:
                        raise PluginError("INPUT_CHANGED", f"Excel 已发生变化或行已执行：{missing_rows}")
                    selected = [by_row[row] for row in requested_rows]
                else:
                    if len(screenshots) != len(items) and not params.get("include_unmatched", False):
                        raise PluginError("SCREENSHOT_COUNT_MISMATCH", f"待执行项 {len(items)} 条，截图 {len(screenshots)} 张，请保持数量一致")
                    selected = items[:len(screenshots)]
                if not selected:
                    raise PluginError("INPUT_INVALID", "没有可关联的待执行项或截图")
                for item, screenshot in zip(selected, screenshots):
                    commit(item, screenshot)

            saved_items = [item for item, _ in saved]
            saved_rows = {item.row_index for item in saved_items}
            skipped_rows = {item.row_index for item in skipped}
            unmatched = [
                {"row_index": item.row_index, "case_name": item.case_name, "checkpoint": item.checkpoint, "status": "skipped" if item.row_index in skipped_rows else "unfinished" if interactive and ended else "missing_screenshot"}
                for item in items if item.row_index not in saved_rows
            ]
            summary = {
                "input_file": cases.path.name,
                **cases.preview_data(),
                "mode": mode,
                "case_count": len({item.case_name for item in saved_items}),
                "evidence_count": len(saved),
                "pending_count": len(unmatched),
                "rows": [item.row_index for item in saved_items],
                "items": [
                    {"row_index": item.row_index, "case_name": item.case_name, "checkpoint": item.checkpoint, "step_name": item.step_name, "status": "saved", "screenshot": f"screenshots/{shot.name}" if interactive else shot.name}
                    for item, shot in saved
                ],
                "reports": reports,
                "excel_updated": update_excel and bool(saved),
                "ended_early": ended,
                "unmatched": unmatched,
            }
            self.context.files.write_text("evidence-index.json", json.dumps(summary, ensure_ascii=False, indent=2))
            files = screenshot_files + reports + ([excel_relative] if update_excel and saved else []) + ["evidence-index.json"]
            warnings = [f"尚有 {len(unmatched)} 条未生成证据" ] if unmatched else []
            message = f"已生成 {len(reports)} 份报告，写入 {len(saved)} 张截图"
            return Result("success", message, summary, files, warnings)
        finally:
            cases.close()

    def destroy(self):
        pass
