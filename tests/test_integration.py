from __future__ import annotations

import csv, importlib.util, json, os, shutil, subprocess, sys, tempfile, threading, time, unittest
from unittest.mock import patch
from datetime import date, timedelta
from pathlib import Path

from testbox.core.errors import ErrorCode
from testbox.core.plugin_packages import install_plugin, package_plugin, uninstall_plugin
from testbox.core.manifest import Manifest
from testbox.core.plugin_registry import PluginManager
from testbox.core.process_runner import HostExecution
from testbox.core.runtime import PluginExecutionLock, Runtime

ROOT = Path(__file__).resolve().parents[1]
MINIMAL_FIELDS = [{"name": "value", "type": "VARCHAR(16)", "generator": "template", "options": {"value": "test-{index}"}}]
PHONE_FIELDS = [{"name": "phone", "type": "VARCHAR(11)", "generator": "mobile_cn"}]

class RuntimeIntegrationTests(unittest.TestCase):
    def setUp(self):
        self.temp = Path(tempfile.mkdtemp()); shutil.copytree(ROOT / "plugins", self.temp / "plugins")
        self.runtime = Runtime(self.temp)
    def tearDown(self):
        self.runtime.close()
        shutil.rmtree(self.temp)
    def test_discovers_bundled_commands(self): self.assertEqual(set(self.runtime.manager.available), {"data.mock", "sql.parse", "sql.select", "evidence.build"})
    def test_evidence_declares_non_concurrent_execution(self):
        self.assertFalse(self.runtime.manager.available["evidence.build"].capabilities["concurrency"])
    def test_plugin_execution_lock_serializes_callers(self):
        lock_path = self.temp / "locks" / "evidence.lock"; first = PluginExecutionLock(lock_path); second = PluginExecutionLock(lock_path)
        first.acquire(); acquired = []; waiting = threading.Event()

        def acquire_second():
            waiting.set(); second.acquire(); acquired.append(True); second.release()

        thread = threading.Thread(target=acquire_second); thread.start()
        self.assertTrue(waiting.wait(timeout=1)); time.sleep(0.05); self.assertEqual(acquired, [])
        first.release(); thread.join(timeout=2)
        self.assertFalse(thread.is_alive()); self.assertEqual(acquired, [True])

    def test_plugin_execution_lock_serializes_processes(self):
        lock_path = self.temp / "locks" / "evidence.lock"; released = self.temp / "released.marker"
        script = "from pathlib import Path; from testbox.core.runtime import PluginExecutionLock; import sys,time; lock=PluginExecutionLock(Path(sys.argv[1])); lock.acquire(); print('locked', flush=True); time.sleep(0.25); lock.release(); Path(sys.argv[2]).write_text('released')"
        process = subprocess.Popen([sys.executable, "-c", script, str(lock_path), str(released)], stdout=subprocess.PIPE, text=True)
        self.assertEqual(process.stdout.readline().strip(), "locked")
        second = PluginExecutionLock(lock_path); second.acquire()
        process.wait(timeout=2)
        self.assertTrue(released.is_file()); second.release()
        process.stdout.close()
    def test_concurrent_runtime_does_not_abandon_started_host(self):
        outer = self

        class InProcessRunner:
            def run(self, request, *, task_id, on_started=None):
                outer.assertIsNotNone(on_started)
                on_started(os.getpid())
                # A separate CLI/GUI Runtime can start while this task is
                # running. Startup recovery must see the live Host PID and
                # leave the task in RUNNING state.
                peer = Runtime(outer.temp)
                try:
                    outer.assertEqual(peer.get_task(task_id)["status"], "RUNNING")
                finally:
                    peer.close()
                output = Path(request["workspace"]) / "output" / f"{task_id}.json"
                output.write_text("[]", encoding="utf-8")
                return HostExecution(
                    {"status": "success", "message": "ok", "data": {}, "files": [f"{task_id}.json"], "warnings": []},
                    0,
                    "",
                    os.getpid(),
                )

        self.runtime.process_runner = InProcessRunner()
        task_id, result = self.runtime.run("data.mock", {"count": 1, "format": "json", "fields": MINIMAL_FIELDS})
        self.assertEqual(result.status, "success")
        self.assertEqual(self.runtime.get_task(task_id)["status"], "SUCCEEDED")

    def test_runtime_falls_back_to_source_root_outside_checkout(self):
        working_directory = self.temp / "empty-working-directory"
        working_directory.mkdir()
        environment = os.environ.copy()
        environment["PYTHONPATH"] = str(ROOT) + os.pathsep + environment.get("PYTHONPATH", "")
        process = subprocess.run(
            [sys.executable, "-c", "from testbox.core.runtime import Runtime; r=Runtime(); print(sorted(r.list_commands())); r.close()"],
            cwd=self.temp / "empty-working-directory",
            capture_output=True,
            text=True,
            timeout=10,
            env=environment,
        )
        self.assertEqual(process.returncode, 0, process.stderr)
        self.assertIn("data.mock", process.stdout)

    def test_cli_json_task_commands_and_export(self):
        environment = os.environ.copy()
        environment["PYTHONPATH"] = str(ROOT) + os.pathsep + environment.get("PYTHONPATH", "")

        def cli(*args):
            return subprocess.run([sys.executable, "-m", "testbox.cli", *args], cwd=self.temp, capture_output=True, text=True, timeout=20, env=environment)

        inspect = cli("--json", "plugin", "inspect", "data-generator")
        self.assertEqual(inspect.returncode, 0, inspect.stderr)
        self.assertEqual(json.loads(inspect.stdout)["name"], "data-generator")
        run = cli("--json", "run", "data.mock", "--set", "count=1", "--set", "format=json", "--set", "seed=7", "--set", 'fields=[{"name":"value","generator":"template","options":{"value":"test-{index}"}}]')
        self.assertEqual(run.returncode, 0, run.stderr)
        task = json.loads(run.stdout)
        task_id = task["task_id"]
        listed = cli("--json", "task", "list", "--command", "data.mock", "--limit", "1")
        self.assertEqual(listed.returncode, 0, listed.stderr)
        self.assertEqual(listed and json.loads(listed.stdout)["tasks"][0]["id"], task_id)
        result = cli("--json", "task", "result", task_id)
        self.assertEqual(result.returncode, 0, result.stderr)
        self.assertEqual(json.loads(result.stdout)["status"], "success")
        destination = self.temp / "exported.json"
        exported = cli("--json", "task", "export", task_id, f"{task_id}.json", "--output", str(destination))
        self.assertEqual(exported.returncode, 0, exported.stderr)
        self.assertTrue(destination.is_file())

    def test_cli_uses_stable_exit_code_and_json_error(self):
        environment = os.environ.copy()
        environment["PYTHONPATH"] = str(ROOT) + os.pathsep + environment.get("PYTHONPATH", "")
        process = subprocess.run([sys.executable, "-m", "testbox.cli", "--json", "run", "unknown.command"], cwd=self.temp, capture_output=True, text=True, timeout=20, env=environment)
        self.assertEqual(process.returncode, 2)
        self.assertEqual(json.loads(process.stderr)["error"]["code"], "COMMAND_NOT_FOUND")

    def test_cli_listing_survives_legacy_console_encoding(self):
        environment = os.environ.copy(); environment["PYTHONIOENCODING"] = "cp1252"; environment["PYTHONPATH"] = str(ROOT) + os.pathsep + environment.get("PYTHONPATH", "")
        process = subprocess.run([sys.executable, "-m", "testbox.cli", "plugin", "list"], cwd=self.temp, capture_output=True, text=True, timeout=10, env=environment)
        self.assertEqual(process.returncode, 0, process.stderr)
        self.assertIn("data.mock\tdata-generator", process.stdout)
    def test_gui_entrypoint_can_run_plugin_host_without_desktop_dependencies(self):
        workspace = self.temp / "gui-host-workspace"
        for child in ("input", "output", "logs"):
            (workspace / child).mkdir(parents=True, exist_ok=True)
        request = {"protocol_version": 1, "task_id": "gui-host-test", "plugin_path": str(self.temp / "plugins" / "data-generator"), "entry": "src.main:Plugin", "command": "data.mock", "params": {"count": 1, "format": "json", "seed": 7, "fields": MINIMAL_FIELDS}, "config": {}, "workspace": str(workspace)}
        process = subprocess.run([sys.executable, "-m", "testbox.gui", "--plugin-host"], input=json.dumps(request), capture_output=True, text=True, timeout=10)
        self.assertEqual(process.returncode, 0, process.stderr)
        # The host protocol must remain portable when Windows uses a legacy
        # subprocess-pipe code page; JSON escaping preserves Unicode semantics.
        self.assertTrue(process.stdout.isascii())
        event = json.loads(process.stdout)
        self.assertEqual(event["result"]["status"], "success")
        self.assertTrue((workspace / "output" / event["result"]["files"][0]).is_file())
    def test_mock_is_repeatable_and_traced(self):
        first_id, first = self.runtime.run("data.mock", {"count": 2, "format": "json", "seed": 7, "fields": MINIMAL_FIELDS})
        second_id, second = self.runtime.run("data.mock", {"count": 2, "format": "json", "seed": 7, "fields": MINIMAL_FIELDS})
        self.assertEqual(first.status, "success"); self.assertEqual(second.status, "success")
        self.assertEqual(first.files, [f"{first_id}.json"]); self.assertEqual(second.files, [f"{second_id}.json"])
        first_data = (self.temp / "workspace" / first_id / "output" / first.files[0]).read_text()
        second_data = (self.temp / "workspace" / second_id / "output" / second.files[0]).read_text()
        self.assertEqual(first_data, second_data); self.assertTrue((self.temp / "workspace" / first_id / "result.json").exists()); self.assertTrue((self.temp / "workspace" / first_id / "report.md").exists())
    def test_sql_parser(self):
        sql = self.temp / "schema.sql"; sql.write_text("CREATE TABLE users (id INT NOT NULL COMMENT 'ID', name VARCHAR(50) COMMENT '姓名');", encoding="utf-8")
        task_id, result = self.runtime.run("sql.parse", {"input": str(sql), "format": "json"})
        self.assertEqual(result.status, "success"); self.assertEqual(result.data["field_count"], 2)
        self.assertEqual(result.files, [f"{task_id}.json"])
        self.assertEqual(result.data["output_file"], f"{task_id}.json")
        manifest = json.loads((self.temp / "workspace" / task_id / "manifest.json").read_text(encoding="utf-8"))
        self.assertEqual(manifest["inputs"][0]["parameter"], "input")
        self.assertTrue((self.temp / "workspace" / task_id / manifest["inputs"][0]["staged_path"]).is_file())
        self.assertEqual(manifest["inputs"][0]["sha256"], __import__("hashlib").sha256(sql.read_bytes()).hexdigest())
    def test_evidence_interactive_is_default(self):
        schema = self.runtime.get_command_schema("evidence.build")
        self.assertTrue(schema["properties"]["interactive"]["default"])

    def test_evidence_interactive_mode_owns_screenshot_capture(self):
        cases = self.temp / "mode.xlsx"
        cases.write_bytes(b"placeholder")
        shot = self.temp / "mode.png"
        shot.write_bytes(b"placeholder")
        _, result = self.runtime.run("evidence.build", {"input": str(cases), "screenshots": [str(shot)]})
        self.assertEqual(result.status, "failed")
        self.assertIn("不接收 screenshots", result.message)

    def test_evidence_batch_mode_requires_explicit_screenshots(self):
        cases = self.temp / "mode.xlsx"
        cases.write_bytes(b"placeholder")
        _, result = self.runtime.run("evidence.build", {"input": str(cases), "interactive": False})
        self.assertEqual(result.status, "failed")
        self.assertIn("必须提供 screenshots", result.message)

    def test_evidence_build_discovers_and_builds(self):
        try:
            from openpyxl import Workbook, load_workbook
            from PIL import Image
            from docx import Document
        except ModuleNotFoundError:
            self.skipTest("evidence-tool optional dependencies are not installed")
        cases = self.temp / "cases.xlsx"; workbook = Workbook(); sheet = workbook.active
        sheet.append(["Case Name", "Check Point", "Step", "Desc", "Expected", "Status"])
        sheet.append(["登录/功能", "登录成功", "1", "输入账号密码", "进入首页", ""]); workbook.save(cases)
        shot = self.temp / "shot.png"; Image.new("RGB", (320, 180), "white").save(shot)
        task_id, result = self.runtime.run("evidence.build", {"input": str(cases), "interactive": False, "screenshots": [str(shot)], "update_excel": True})
        self.assertEqual(result.status, "success"); self.assertEqual(result.data["evidence_count"], 1)
        self.assertEqual(result.data["header_row"], 1)
        self.assertEqual(result.data["pending_count"], 0)
        self.assertEqual(result.data["mapping"]["测试名称"], "Case Name")
        output = self.temp / "workspace" / task_id / "output"
        report = next((output / "reports").glob("登录_功能-*.docx")); self.assertTrue(report.exists())
        self.assertIn("验证点：登录成功", [paragraph.text for paragraph in Document(report).paragraphs])
        updated = load_workbook(output / "executed-cases.xlsx"); self.assertEqual(updated.active["F2"].value, "已执行"); updated.close()
        manifest = json.loads((self.temp / "workspace" / task_id / "manifest.json").read_text(encoding="utf-8"))
        self.assertEqual([record["parameter"] for record in manifest["inputs"]], ["input", "screenshots"])
    def test_evidence_handles_large_merged_partial_workflow_and_commits_excel(self):
        try:
            from openpyxl import Workbook, load_workbook
            from PIL import Image
        except ModuleNotFoundError:
            self.skipTest("evidence-tool optional dependencies are not installed")
        cases = self.temp / "many-cases.xlsx"; workbook = Workbook(); sheet = workbook.active
        sheet.append(["说明"]); sheet.append(["测试名称", "验证点", "步骤名称", "步骤描述", "预期结果", "测试结果"])
        for index in range(25): sheet.append([f"用例{index + 1}", f"验证{index + 1}", str(index + 1), "执行", "成功", ""])
        workbook.save(cases)
        shot = self.temp / "partial.png"; Image.new("RGB", (320, 180), "white").save(shot)
        selected_row = 24
        mapping = {"测试名称": "测试名称", "验证点": "验证点", "步骤名称": "步骤名称", "步骤描述": "步骤描述", "预期结果": "预期结果", "测试结果": "测试结果"}
        task_id, result = self.runtime.run("evidence.build", {"input": str(cases), "screenshots": [str(shot)], "interactive": False, "row_indexes": [selected_row], "column_mapping": mapping, "update_excel": True})
        self.assertEqual(result.status, "success"); self.assertEqual(result.data["rows"], [selected_row])
        excel_output = next(name for name in result.files if name.startswith("executed-"))
        self.runtime.commit_output(task_id, excel_output, cases)
        updated = load_workbook(cases); self.assertEqual(updated.active.cell(selected_row, 6).value, "已执行"); self.assertIsNone(updated.active.cell(3, 6).value); updated.close()
        self.assertEqual(result.data["pending_count"], 24)
    def test_evidence_disambiguates_sanitized_report_names(self):
        try:
            from openpyxl import Workbook
            from PIL import Image
        except ModuleNotFoundError:
            self.skipTest("evidence-tool optional dependencies are not installed")
        cases = self.temp / "collisions.xlsx"; workbook = Workbook(); sheet = workbook.active
        sheet.append(["测试名称", "验证点", "测试结果"]); sheet.append(["登录/功能", "A", ""]); sheet.append(["登录:功能", "B", ""]); workbook.save(cases)
        shots = []
        for index in range(2):
            path = self.temp / f"collision-{index}.png"; Image.new("RGB", (100, 60), "white").save(path); shots.append(str(path))
        task_id, result = self.runtime.run("evidence.build", {"input": str(cases), "interactive": False, "screenshots": shots, "row_indexes": [2, 3]})
        reports = [name for name in result.files if name.endswith(".docx")]
        self.assertEqual(len(reports), 2); self.assertEqual(len(set(reports)), 2)
        self.assertTrue(all((self.temp / "workspace" / task_id / "output" / name).is_file() for name in reports))
    def test_evidence_appends_existing_report_without_duplicate_checkpoint(self):
        try:
            from openpyxl import Workbook
            from PIL import Image
            from docx import Document
        except ModuleNotFoundError:
            self.skipTest("evidence-tool optional dependencies are not installed")
        cases = self.temp / "append.xlsx"; workbook = Workbook(); sheet = workbook.active
        sheet.append(["测试名称", "验证点", "步骤名称", "测试结果"]); sheet.append(["支付功能", "支付成功", "1", ""]); workbook.save(cases)
        shot = self.temp / "append.png"; Image.new("RGB", (100, 60), "white").save(shot)
        first_id, first = self.runtime.run("evidence.build", {"input": str(cases), "interactive": False, "screenshots": [str(shot)], "row_indexes": [2], "update_excel": False})
        first_report = self.temp / "workspace" / first_id / "output" / first.files[0]
        second_id, second = self.runtime.run("evidence.build", {"input": str(cases), "interactive": False, "screenshots": [str(shot)], "row_indexes": [2], "existing_reports": [str(first_report)], "update_excel": False})
        second_report = self.temp / "workspace" / second_id / "output" / second.files[0]; document = Document(second_report)
        self.assertEqual([p.text for p in document.paragraphs].count("验证点：支付成功"), 1)
        self.assertEqual(len(document.inline_shapes), 2)
    def test_evidence_inherits_merged_case_and_checkpoint_per_step(self):
        try:
            from openpyxl import Workbook
            from PIL import Image
        except ModuleNotFoundError:
            self.skipTest("evidence-tool optional dependencies are not installed")
        cases = self.temp / "merged.xlsx"; workbook = Workbook(); sheet = workbook.active
        sheet.append(["测试名称", "验证点", "步骤名称", "步骤描述", "预期结果", "测试结果"])
        sheet.merge_cells("A2:A4"); sheet.merge_cells("B2:B4"); sheet["A2"] = "合并用例"; sheet["B2"] = "合并验证点"
        for row in range(2, 5): sheet.cell(row, 3).value = str(row - 1); sheet.cell(row, 4).value = "执行"; sheet.cell(row, 5).value = "成功"
        sheet["F2"] = "已执行"; workbook.save(cases)
        shot1 = self.temp / "merged-1.png"; shot2 = self.temp / "merged-2.png"
        Image.new("RGB", (100, 60), "white").save(shot1); Image.new("RGB", (100, 60), "white").save(shot2)
        _, result = self.runtime.run("evidence.build", {"input": str(cases), "interactive": False, "screenshots": [str(shot1), str(shot2)], "row_indexes": [3, 4], "update_excel": False})
        self.assertEqual([item["row_index"] for item in result.data["items"]], [3, 4])
        self.assertTrue(all(item["case_name"] == "合并用例" and item["checkpoint"] == "合并验证点" for item in result.data["items"]))
    def test_sql_parser_exports_xlsx(self):
        sql = self.temp / "schema.sql"; sql.write_text("CREATE TABLE users (id INT NOT NULL);", encoding="utf-8")
        task_id, result = self.runtime.run("sql.parse", {"input": str(sql), "format": "xlsx"})
        self.assertEqual(result.status, "success")
        self.assertEqual(result.data["format"], "xlsx")
        self.assertEqual(result.data["output_file"], f"{task_id}.xlsx")
        output = self.temp / "workspace" / task_id / "output" / result.files[0]
        with __import__("zipfile").ZipFile(output) as archive:
            self.assertIn("xl/worksheets/sheet1.xml", archive.namelist())
            self.assertIn("字段清单", archive.read("xl/workbook.xml").decode("utf-8"))
    def test_sql_parser_exports_csv_with_same_field_data(self):
        sql = self.temp / "schema.sql"
        sql.write_text("CREATE TABLE users (id INT NOT NULL COMMENT '用户ID', name VARCHAR(50) COMMENT '姓名');", encoding="utf-8")
        task_id, result = self.runtime.run("sql.parse", {"input": str(sql), "format": "csv"})
        self.assertEqual(result.status, "success")
        self.assertEqual(result.data["format"], "csv")
        self.assertEqual(result.data["output_file"], f"{task_id}.csv")
        output = self.temp / "workspace" / task_id / "output" / result.files[0]
        with output.open(encoding="utf-8", newline="") as csv_file:
            rows = list(csv.DictReader(csv_file))
        self.assertEqual([row["field"] for row in rows], ["id", "name"])
        self.assertEqual(rows[0]["comment"], "用户ID")
        self.assertEqual(rows[1]["length"], "50")

    def test_sql_select_consumes_sql_parser_json_output(self):
        sql = self.temp / "schema.sql"
        sql.write_text("CREATE TABLE users (id INT NOT NULL COMMENT '用户ID', name VARCHAR(50) COMMENT '姓名'); CREATE TABLE audit_log (event_id BIGINT);", encoding="utf-8")
        parse_task, parsed = self.runtime.run("sql.parse", {"input": str(sql), "format": "json"})
        self.assertEqual(parsed.status, "success")
        parse_output = self.temp / "workspace" / parse_task / "output" / parsed.files[0]
        self.assertEqual(parsed.files, [f"{parse_task}.json"])

        select_task, selected = self.runtime.run("sql.select", {"input": str(parse_output), "dialect": "mysql", "include_comments": True})
        self.assertEqual(selected.status, "success")
        self.assertEqual(selected.files, [f"{select_task}.sql"])
        output = self.temp / "workspace" / select_task / "output" / selected.files[0]
        content = output.read_text(encoding="utf-8")
        self.assertIn("SELECT", content)
        self.assertIn("    `id` /* 用户ID */", content)
        self.assertIn("FROM `users`;", content)
        self.assertIn("FROM `audit_log`;", content)
        self.assertEqual(selected.data["table_count"], 2)

    def test_sql_select_consumes_sql_parser_xlsx_output(self):
        sql = self.temp / "schema.sql"
        sql.write_text("CREATE TABLE users (id INT, name VARCHAR(50));", encoding="utf-8")
        parse_task, parsed = self.runtime.run("sql.parse", {"input": str(sql), "format": "xlsx"})
        parse_output = self.temp / "workspace" / parse_task / "output" / parsed.files[0]
        select_task, selected = self.runtime.run("sql.select", {"input": str(parse_output), "dialect": "postgresql"})
        self.assertEqual(selected.status, "success")
        content = (self.temp / "workspace" / select_task / "output" / selected.files[0]).read_text(encoding="utf-8")
        self.assertIn('FROM "users";', content)
        self.assertIn('"name"', content)

    def test_sql_parser_supports_nested_types_constraints_and_postgres_comments(self):
        sql = self.temp / "mixed.sql"
        sql.write_text("""CREATE TABLE public.orders (
            id BIGINT GENERATED ALWAYS AS IDENTITY,
            customer_id UUID NOT NULL,
            amount DECIMAL(18,2) DEFAULT 0.00,
            metadata JSONB,
            CONSTRAINT orders_pk PRIMARY KEY (id),
            CONSTRAINT orders_customer_fk FOREIGN KEY (customer_id) REFERENCES customer(id)
        );
        COMMENT ON TABLE public.orders IS '订单表';
        COMMENT ON COLUMN public.orders.amount IS '交易金额';
        """, encoding="utf-8")
        task_id, result = self.runtime.run("sql.parse", {"input": str(sql), "format": "json", "dialect": "postgresql"})
        rows = json.loads((self.temp / "workspace" / task_id / "output" / result.files[0]).read_text(encoding="utf-8"))
        by_name = {row["field"]: row for row in rows}
        self.assertEqual(by_name["amount"]["type"], "DECIMAL(18,2)")
        self.assertEqual(by_name["amount"]["comment"], "交易金额")
        self.assertTrue(by_name["id"]["primary_key"])
        self.assertEqual(by_name["customer_id"]["foreign_table"], "customer")
        self.assertTrue(by_name["id"]["auto_increment"])
    def test_sql_parser_supports_sqlserver_and_fail_on_warnings(self):
        sql = self.temp / "sqlserver.sql"
        sql.write_text("CREATE TABLE [dbo].[users] ([id] INT IDENTITY(1,1) NOT NULL PRIMARY KEY, [name] NVARCHAR(100) NULL);", encoding="utf-8")
        task_id, result = self.runtime.run("sql.parse", {"input": str(sql), "format": "json", "dialect": "sqlserver"})
        rows = json.loads((self.temp / "workspace" / task_id / "output" / result.files[0]).read_text(encoding="utf-8"))
        self.assertEqual(rows[0]["table"], "dbo.users")
        self.assertTrue(rows[0]["auto_increment"])
    def test_sql_parser_dimensions_partition_literals_and_constraint_toggle(self):
        sql = self.temp / "partitioned.sql"
        sql.write_text("""CREATE TABLE `sales` (
            `id` BIGINT NOT NULL,
            `code` VARCHAR(32) DEFAULT 'A,--B',
            `amount` DECIMAL(18, 2) DEFAULT 0,
            `note` TEXT DEFAULT 'it''s; valid',
            PRIMARY KEY (`id`),
            UNIQUE (`code`)
        ) PARTITION BY HASH (`id`) PARTITIONS 4;""", encoding="utf-8")
        task_id, result = self.runtime.run("sql.parse", {"input": str(sql), "format": "json", "include_constraints": False})
        rows = json.loads((self.temp / "workspace" / task_id / "output" / result.files[0]).read_text(encoding="utf-8"))
        by_name = {row["field"]: row for row in rows}
        self.assertEqual(by_name["code"]["length"], 32)
        self.assertEqual((by_name["amount"]["precision"], by_name["amount"]["scale"]), (18, 2))
        self.assertEqual(by_name["code"]["default"], "'A,--B'")
        self.assertEqual(by_name["note"]["default"], "'it''s; valid'")
        self.assertEqual(by_name["id"]["partition"], "HASH (`id`)")
        self.assertNotIn("primary_key", by_name["id"])
        self.assertNotIn("foreign_table", by_name["id"])
    def test_sql_parser_mysql_common_ddl(self):
        sql = self.temp / "mysql.sql"
        sql.write_text("""CREATE TABLE `shop`.`orders` (
            `id` BIGINT UNSIGNED NOT NULL AUTO_INCREMENT COMMENT '订单ID',
            `status` ENUM('NEW','PAID','CLOSED') NOT NULL DEFAULT 'NEW',
            `updated_at` TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
            `amount` DECIMAL(20,4) UNSIGNED DEFAULT 0.0000,
            PRIMARY KEY (`id`),
            UNIQUE KEY `uk_status` (`status`)
        ) COMMENT='订单表' PARTITION BY HASH (`id`) PARTITIONS 8;""", encoding="utf-8")
        task_id, result = self.runtime.run("sql.parse", {"input": str(sql), "format": "json", "dialect": "mysql"})
        self.assertEqual(result.status, "success")
        rows = json.loads((self.temp / "workspace" / task_id / "output" / result.files[0]).read_text(encoding="utf-8"))
        by_name = {row["field"]: row for row in rows}
        self.assertEqual(by_name["id"]["type"], "BIGINT UNSIGNED")
        self.assertTrue(by_name["id"]["auto_increment"])
        self.assertEqual(by_name["status"]["type"], "ENUM('NEW','PAID','CLOSED')")
        self.assertTrue(by_name["status"]["unique"])
        self.assertEqual(by_name["updated_at"]["default"], "CURRENT_TIMESTAMP")
        self.assertEqual(by_name["amount"]["precision"], 20)
        self.assertEqual(by_name["amount"]["scale"], 4)
        self.assertEqual(by_name["id"]["table_comment"], "订单表")
        self.assertEqual(by_name["id"]["partition"], "HASH (`id`)")
    def test_sql_parser_postgresql_common_ddl(self):
        sql = self.temp / "postgresql.sql"
        sql.write_text("""CREATE TABLE public.accounts (
            id BIGSERIAL PRIMARY KEY,
            owner_id UUID NOT NULL,
            balance NUMERIC(19,2) DEFAULT 0::numeric,
            tags TEXT[] DEFAULT ARRAY[]::TEXT[],
            created_at TIMESTAMP WITH TIME ZONE DEFAULT now(),
            CONSTRAINT accounts_owner_fk FOREIGN KEY (owner_id) REFERENCES auth.users(id)
        );
        COMMENT ON TABLE public.accounts IS '账户表';
        COMMENT ON COLUMN public.accounts.balance IS '账面''余额';""", encoding="utf-8")
        task_id, result = self.runtime.run("sql.parse", {"input": str(sql), "format": "json", "dialect": "postgresql"})
        self.assertEqual(result.status, "success")
        rows = json.loads((self.temp / "workspace" / task_id / "output" / result.files[0]).read_text(encoding="utf-8"))
        by_name = {row["field"]: row for row in rows}
        self.assertTrue(by_name["id"]["auto_increment"])
        self.assertEqual(by_name["balance"]["comment"], "账面'余额")
        self.assertEqual(by_name["balance"]["default"], "0::numeric")
        self.assertEqual(by_name["tags"]["type"], "TEXT[]")
        self.assertEqual(by_name["owner_id"]["foreign_table"], "auth.users")
    def test_sql_parser_sqlserver_common_ddl(self):
        sql = self.temp / "sqlserver.sql"
        sql.write_text("""CREATE TABLE [sales].[invoice] (
            [invoice_id] BIGINT IDENTITY(1000,1) NOT NULL,
            [customer_id] UNIQUEIDENTIFIER NOT NULL,
            [total] DECIMAL(18,2) CONSTRAINT [DF_invoice_total] DEFAULT ((0)),
            [created_at] DATETIME2(3) DEFAULT (SYSUTCDATETIME()),
            CONSTRAINT [PK_invoice] PRIMARY KEY ([invoice_id]),
            CONSTRAINT [FK_invoice_customer] FOREIGN KEY ([customer_id]) REFERENCES [crm].[customer]([id])
        );""", encoding="utf-8")
        task_id, result = self.runtime.run("sql.parse", {"input": str(sql), "format": "json", "dialect": "sqlserver"})
        self.assertEqual(result.status, "success")
        rows = json.loads((self.temp / "workspace" / task_id / "output" / result.files[0]).read_text(encoding="utf-8"))
        by_name = {row["field"]: row for row in rows}
        self.assertTrue(by_name["invoice_id"]["primary_key"])
        self.assertTrue(by_name["invoice_id"]["auto_increment"])
        self.assertEqual(by_name["total"]["default"], "((0))")
        self.assertEqual(by_name["created_at"]["length"], 3)
        self.assertEqual(by_name["customer_id"]["foreign_table"], "crm.customer")
    def test_sql_parser_oracle_common_ddl(self):
        sql = self.temp / "oracle.sql"
        sql.write_text("""CREATE TABLE APP.CUSTOMER (
            ID NUMBER(12) GENERATED BY DEFAULT AS IDENTITY,
            NAME VARCHAR2(100 CHAR) NOT NULL,
            CREDIT_LIMIT NUMBER(15,2) DEFAULT 0,
            CREATED_AT TIMESTAMP(6) DEFAULT SYSTIMESTAMP,
            CONSTRAINT PK_CUSTOMER PRIMARY KEY (ID)
        );
        COMMENT ON TABLE APP.CUSTOMER IS '客户表';
        COMMENT ON COLUMN APP.CUSTOMER.NAME IS '客户名称';""", encoding="utf-8")
        task_id, result = self.runtime.run("sql.parse", {"input": str(sql), "format": "json", "dialect": "oracle"})
        self.assertEqual(result.status, "success")
        rows = json.loads((self.temp / "workspace" / task_id / "output" / result.files[0]).read_text(encoding="utf-8"))
        by_name = {row["field"]: row for row in rows}
        self.assertTrue(by_name["ID"]["auto_increment"])
        self.assertEqual(by_name["ID"]["precision"], 12)
        self.assertEqual(by_name["NAME"]["length"], 100)
        self.assertEqual(by_name["NAME"]["comment"], "客户名称")
    def test_sql_parser_sqlite_common_ddl(self):
        sql = self.temp / "sqlite.sql"
        sql.write_text("""CREATE TABLE IF NOT EXISTS "order,item" (
            "id" INTEGER PRIMARY KEY AUTOINCREMENT,
            "code,value" TEXT NOT NULL UNIQUE,
            "enabled" INTEGER NOT NULL DEFAULT 1,
            "payload" BLOB,
            "parent_id" INTEGER REFERENCES "order,item"("id")
        );""", encoding="utf-8")
        task_id, result = self.runtime.run("sql.parse", {"input": str(sql), "format": "json", "dialect": "sqlite"})
        self.assertEqual(result.status, "success")
        rows = json.loads((self.temp / "workspace" / task_id / "output" / result.files[0]).read_text(encoding="utf-8"))
        by_name = {row["field"]: row for row in rows}
        self.assertEqual(rows[0]["table"], "order,item")
        self.assertTrue(by_name["id"]["auto_increment"])
        self.assertTrue(by_name["code,value"]["unique"])
        self.assertEqual(by_name["parent_id"]["foreign_table"], "order,item")
    def test_sql_parser_hudi_spark_ddl(self):
        sql = self.temp / "hudi.sql"
        sql.write_text("""CREATE TABLE lakehouse.hudi_orders (
            order_id BIGINT COMMENT '订单ID',
            customer STRUCT<id:BIGINT,name:STRING>,
            items ARRAY<STRUCT<sku:STRING,qty:INT,price:DECIMAL(18,2)>>,
            attributes MAP<STRING,STRING>,
            updated_at TIMESTAMP,
            dt STRING
        ) USING HUDI
        PARTITIONED BY (dt)
        TBLPROPERTIES (
            'primaryKey' = 'order_id',
            'preCombineField' = 'updated_at'
        );""", encoding="utf-8")
        task_id, result = self.runtime.run("sql.parse", {"input": str(sql), "format": "json", "dialect": "hudi"})
        self.assertEqual(result.status, "success")
        rows = json.loads((self.temp / "workspace" / task_id / "output" / result.files[0]).read_text(encoding="utf-8"))
        by_name = {row["field"]: row for row in rows}
        self.assertEqual(len(rows), 6)
        self.assertEqual(by_name["customer"]["type"], "STRUCT<ID:BIGINT,NAME:STRING>")
        self.assertEqual(by_name["items"]["type"], "ARRAY<STRUCT<SKU:STRING,QTY:INT,PRICE:DECIMAL(18,2)>>")
        self.assertEqual(by_name["attributes"]["type"], "MAP<STRING,STRING>")
        self.assertEqual(by_name["order_id"]["partition"], "(dt)")
    def test_sql_parser_hive_ddl(self):
        sql = self.temp / "hive.sql"
        sql.write_text("""CREATE TABLE IF NOT EXISTS ods.user_events (
            user_id BIGINT COMMENT '用户ID',
            event_name STRING,
            properties MAP<STRING,STRING>,
            labels ARRAY<STRING>,
            profile STRUCT<age:INT,city:STRING>
        ) COMMENT '用户行为明细'
        PARTITIONED BY (dt STRING COMMENT '业务日期', hour STRING)
        STORED AS PARQUET
        TBLPROPERTIES ('parquet.compression'='SNAPPY');""", encoding="utf-8")
        task_id, result = self.runtime.run("sql.parse", {"input": str(sql), "format": "json", "dialect": "hive"})
        self.assertEqual(result.status, "success")
        rows = json.loads((self.temp / "workspace" / task_id / "output" / result.files[0]).read_text(encoding="utf-8"))
        by_name = {row["field"]: row for row in rows}
        self.assertEqual(len(rows), 5)
        self.assertEqual(by_name["properties"]["type"], "MAP<STRING,STRING>")
        self.assertEqual(by_name["profile"]["type"], "STRUCT<AGE:INT,CITY:STRING>")
        self.assertEqual(by_name["user_id"]["table_comment"], "用户行为明细")
        self.assertEqual(by_name["user_id"]["partition"], "(dt STRING COMMENT '业务日期', hour STRING)")
    def test_sql_parser_hbase_phoenix_ddl(self):
        sql = self.temp / "hbase-phoenix.sql"
        sql.write_text("""CREATE TABLE IF NOT EXISTS APP.CUSTOMER (
            TENANT_ID VARCHAR NOT NULL,
            CUSTOMER_ID BIGINT NOT NULL,
            INFO.NAME VARCHAR,
            INFO.CREATED_AT UNSIGNED_TIMESTAMP,
            METRIC.BALANCE DECIMAL(18,2),
            CONSTRAINT PK_CUSTOMER PRIMARY KEY (TENANT_ID, CUSTOMER_ID)
        ) SALT_BUCKETS=8, VERSIONS=3, COMPRESSION='SNAPPY';""", encoding="utf-8")
        task_id, result = self.runtime.run("sql.parse", {"input": str(sql), "format": "json", "dialect": "hbase"})
        self.assertEqual(result.status, "success")
        rows = json.loads((self.temp / "workspace" / task_id / "output" / result.files[0]).read_text(encoding="utf-8"))
        by_name = {row["field"]: row for row in rows}
        self.assertEqual(len(rows), 5)
        self.assertTrue(by_name["TENANT_ID"]["primary_key"])
        self.assertTrue(by_name["CUSTOMER_ID"]["primary_key"])
        self.assertEqual(by_name["INFO.CREATED_AT"]["type"], "UNSIGNED_TIMESTAMP")
        self.assertEqual(by_name["METRIC.BALANCE"]["precision"], 18)
    def test_sql_parser_maxcompute_mc_ddl(self):
        sql = self.temp / "maxcompute.sql"
        sql.write_text("""CREATE TABLE IF NOT EXISTS mart.order_detail (
            order_id BIGINT COMMENT '订单ID',
            buyer_id STRING COMMENT '买家ID',
            amount DECIMAL(20,4),
            tags ARRAY<STRING>,
            ext MAP<STRING,STRING>,
            item STRUCT<sku:STRING,qty:BIGINT>
        ) COMMENT '订单明细'
        PARTITIONED BY (ds STRING COMMENT '分区日期')
        LIFECYCLE 365;""", encoding="utf-8")
        task_id, result = self.runtime.run("sql.parse", {"input": str(sql), "format": "json", "dialect": "mc"})
        self.assertEqual(result.status, "success")
        rows = json.loads((self.temp / "workspace" / task_id / "output" / result.files[0]).read_text(encoding="utf-8"))
        by_name = {row["field"]: row for row in rows}
        self.assertEqual(len(rows), 6)
        self.assertEqual((by_name["amount"]["precision"], by_name["amount"]["scale"]), (20, 4))
        self.assertEqual(by_name["tags"]["type"], "ARRAY<STRING>")
        self.assertEqual(by_name["item"]["type"], "STRUCT<SKU:STRING,QTY:BIGINT>")
        self.assertEqual(by_name["order_id"]["table_comment"], "订单明细")
        self.assertEqual(by_name["order_id"]["partition"], "(ds STRING COMMENT '分区日期')")
    def test_sql_parser_auto_detects_dialects_and_allows_override(self):
        samples = {
            "mysql": "CREATE TABLE t (id BIGINT AUTO_INCREMENT);",
            "postgresql": "CREATE TABLE t (id BIGSERIAL, payload JSONB);",
            "sqlserver": "CREATE TABLE [t] ([id] BIGINT IDENTITY(1,1));",
            "oracle": "CREATE TABLE t (id NUMBER(12), name VARCHAR2(50));",
            "sqlite": "CREATE TABLE t (id INTEGER PRIMARY KEY AUTOINCREMENT);",
            "hudi": "CREATE TABLE t (id BIGINT) USING HUDI TBLPROPERTIES ('primaryKey'='id');",
            "hive": "CREATE TABLE t (id BIGINT) STORED AS PARQUET;",
            "hbase": "CREATE TABLE t (id BIGINT PRIMARY KEY) SALT_BUCKETS=4;",
            "maxcompute": "CREATE TABLE t (id BIGINT) LIFECYCLE 30;",
            "auto": "CREATE TABLE t (id INTEGER, name VARCHAR(20));",
        }
        for expected, ddl in samples.items():
            with self.subTest(expected=expected):
                sql = self.temp / f"auto-{expected}.sql"; sql.write_text(ddl, encoding="utf-8")
                task_id, result = self.runtime.run("sql.parse", {"input": str(sql), "format": "json"})
                self.assertEqual(result.data["requested_dialect"], "auto")
                self.assertEqual(result.data["dialect"], expected)
                rows = json.loads((self.temp / "workspace" / task_id / "output" / result.files[0]).read_text(encoding="utf-8"))
                self.assertTrue(all(row["dialect"] == expected for row in rows))
        sql = self.temp / "override.sql"; sql.write_text(samples["mysql"], encoding="utf-8")
        _, result = self.runtime.run("sql.parse", {"input": str(sql), "format": "json", "dialect": "hive"})
        self.assertEqual(result.data["requested_dialect"], "hive")
        self.assertEqual(result.data["dialect"], "hive")
    def test_plugin_command_conflict_isolated_to_conflicting_plugin(self):
        registry_root = self.temp / "conflict-plugins"
        shutil.copytree(ROOT / "plugins" / "data-generator", registry_root / "first")
        shutil.copytree(ROOT / "plugins" / "data-generator", registry_root / "second")
        for directory, name in ((registry_root / "first", "first-plugin"), (registry_root / "second", "second-plugin")):
            manifest = directory / "manifest.yaml"
            text = manifest.read_text(encoding="utf-8").replace("name: data-generator", f"name: {name}").replace("name: data.mock", "name: conflict.run")
            manifest.write_text(text, encoding="utf-8")
        manager = PluginManager([registry_root])
        manager.discover()
        self.assertIn("conflict.run", manager.available)
        self.assertEqual(len(manager.unavailable), 1)
        self.assertIn("命令冲突", next(iter(manager.unavailable.values())))

    def test_plugin_dependency_failure_is_structured(self):
        plugin = self.temp / "plugins" / "missing-dependency"
        (plugin / "src").mkdir(parents=True)
        (plugin / "manifest.yaml").write_text("""schema_version: 1
name: missing-dependency
version: 1.0.0
description: dependency failure fixture
category: test
core_compatibility: ">=1.0,<2.0"
entry: src.main:Plugin
commands:
  - name: dependency.check
    description: dependency check
capabilities:
  concurrency: true
  network: false
  filesystem: output-only
  resources: []
""", encoding="utf-8")
        (plugin / "src" / "main.py").write_text("""from testbox.sdk import PluginError, Result
class Plugin:
    def init(self, context):
        raise PluginError("DEPENDENCY_MISSING", "missing optional dependency")
    def execute(self, command, params):
        return Result("success", "unexpected")
    def destroy(self):
        pass
""", encoding="utf-8")
        runtime = Runtime(self.temp)
        try:
            _, result = runtime.run("dependency.check", {})
            self.assertEqual(result.status, "failed")
            self.assertEqual(result.data["error_code"], ErrorCode.DEPENDENCY_MISSING)
            self.assertEqual(runtime.history.list_tasks(command="dependency.check", limit=1)[0]["status"], "FAILED")
        finally:
            runtime.close()

    def test_manifest_requires_minimum_capabilities(self):
        plugin = self.temp / "invalid-plugin"; plugin.mkdir()
        (plugin / "manifest.yaml").write_text("""schema_version: 1
name: invalid-plugin
version: 1.0.0
description: invalid
category: test
core_compatibility: \">=1.0,<2.0\"
entry: src.main:Plugin
commands: []
""", encoding="utf-8")
        with self.assertRaisesRegex(ValueError, "capabilities"):
            Manifest.load(plugin / "manifest.yaml")
    def test_manifest_rejects_incompatible_core_version(self):
        plugin = self.temp / "incompatible-plugin"; plugin.mkdir()
        (plugin / "manifest.yaml").write_text("""schema_version: 1
name: incompatible-plugin
version: 1.0.0
description: invalid
category: test
core_compatibility: \">=2.0,<3.0\"
entry: src.main:Plugin
commands:
  - name: test.run
    description: test
capabilities:
  concurrency: true
  network: false
  filesystem: output-only
  resources: []
""", encoding="utf-8")
        with self.assertRaisesRegex(ValueError, "不兼容"):
            Manifest.load(plugin / "manifest.yaml")
    def test_project_config_is_injected_into_plugin(self):
        (self.temp / "config.yaml").write_text('phone_prefixes: "199"\n', encoding="utf-8")
        task_id, result = self.runtime.run("data.mock", {"count": 1, "format": "json", "seed": 7, "fields": PHONE_FIELDS})
        rows = json.loads((self.temp / "workspace" / task_id / "output" / result.files[0]).read_text(encoding="utf-8"))
        self.assertTrue(rows[0]["phone"].startswith("199"))
    def test_schema_maximum_is_rejected(self):
        with self.assertRaisesRegex(ValueError, "不能大于"):
            self.runtime.run("data.mock", {"count": 100001, "format": "json"})
    def test_data_generator_requires_explicit_schema(self):
        _, result = self.runtime.run("data.mock", {"count": 1, "format": "json"})
        self.assertEqual(result.status, "failed")
        self.assertIn("请至少添加一个字段", result.message)

    def test_data_generator_rejects_conflicting_schema_modes(self):
        _, result = self.runtime.run("data.mock", {"count": 1, "format": "json", "fields": MINIMAL_FIELDS, "template": "account"})
        self.assertEqual(result.status, "failed")
        self.assertIn("只能选择一种模式", result.message)

    def test_data_generator_supports_custom_fields_and_types(self):
        fields = [
            {"name": "id", "type": "BIGINT", "generator": "sequence", "unique": True, "options": {"prefix": "ID-", "width": 4}},
            {"name": "status", "type": "VARCHAR(8)", "generator": "weighted_enum", "options": {"values": ["NEW", "DONE"]}},
            {"name": "amount", "type": "DECIMAL(10,2)", "generator": "decimal_random", "options": {"min": 10, "max": 20, "scale": 2}},
            {"name": "enabled", "type": "BOOLEAN", "generator": "boolean_random"},
            {"name": "remark", "type": "VARCHAR(20)", "generator": "constant", "options": {"value": "TEST DATA ONLY"}}
        ]
        task_id, result = self.runtime.run("data.mock", {"count": 3, "format": "json", "seed": 7, "fields": fields})
        rows = json.loads((self.temp / "workspace" / task_id / "output" / result.files[0]).read_text(encoding="utf-8"))
        self.assertEqual(list(rows[0]), ["id", "status", "amount", "enabled", "remark"])
        self.assertEqual([row["id"] for row in rows], ["ID-0001", "ID-0002", "ID-0003"])
        self.assertTrue(all(row["status"] in {"NEW", "DONE"} for row in rows))
        self.assertTrue(all(isinstance(row["amount"], float) for row in rows))
        self.assertTrue(all(row["remark"] == "TEST DATA ONLY" for row in rows))

    def test_data_generator_sql_schema_uses_table_and_constraints(self):
        ddl = self.temp / "orders.sql"
        ddl.write_text("CREATE TABLE orders (id BIGINT PRIMARY KEY, code VARCHAR(5) UNIQUE, active BOOLEAN, amount DECIMAL(8,2));", encoding="utf-8")
        task_id, result = self.runtime.run("data.mock", {"count": 2, "format": "sql", "seed": 7, "source_file": str(ddl), "source_format": "sql"})
        script = (self.temp / "workspace" / task_id / "output" / result.files[0]).read_text(encoding="utf-8")
        self.assertIn('INSERT INTO `orders`', script)
        self.assertIn("`id`", script)
        self.assertIn("`active`", script)
        self.assertEqual(result.data["table"], "orders")

    def test_data_generator_sql_can_include_create_table(self):
        fields = [
            {"name": "id", "type": "BIGINT", "generator": "sequence", "primary_key": True, "unique": True},
            {"name": "status", "type": "VARCHAR(10)", "generator": "constant", "options": {"value": "OK"}},
        ]
        task_id, result = self.runtime.run("data.mock", {"count": 1, "format": "sql", "table": "user_info", "fields": fields, "sql_create_table": True, "sql_transaction": False})
        script = (self.temp / "workspace" / task_id / "output" / result.files[0]).read_text(encoding="utf-8")
        self.assertIn("CREATE TABLE `user_info`", script)
        self.assertIn("`id` BIGINT PRIMARY KEY", script)
        self.assertIn("INSERT INTO `user_info`", script)

    def test_data_generator_supports_field_unique_phone_rule(self):
        rules = [{"name": "mobile", "generator": "mobile_cn", "unique": True, "options": {"prefixes": ["199"]}}]
        task_id, result = self.runtime.run("data.mock", {"count": 100, "format": "json", "seed": 7, "rules": rules})
        rows = json.loads((self.temp / "workspace" / task_id / "output" / result.files[0]).read_text(encoding="utf-8"))
        self.assertEqual(len({row["mobile"] for row in rows}), 100)
        self.assertTrue(all(row["mobile"].startswith("199") and len(row["mobile"]) == 11 for row in rows))
        self.assertEqual(result.data["unique_fields"], ["mobile"])
        self.assertEqual(result.data["administrative_divisions_version"], "2025-12-27-c49d495")
    def test_data_generator_rejects_insufficient_unique_capacity(self):
        rules = [{"name": "status", "generator": "weighted_enum", "unique": True, "options": {"values": ["A", "B"]}}]
        _, result = self.runtime.run("data.mock", {"count": 3, "format": "json", "rules": rules})
        self.assertEqual(result.status, "failed")
        self.assertIn("唯一值容量", result.message)
    def test_data_generator_filters_addresses_by_province_and_city(self):
        rules = [{"name": "address", "generator": "china_address", "options": {"province": ["广东省"], "city": ["深圳市"], "include_virtual_street": False}}]
        task_id, result = self.runtime.run("data.mock", {"count": 3, "format": "json", "seed": 7, "rules": rules})
        rows = json.loads((self.temp / "workspace" / task_id / "output" / result.files[0]).read_text(encoding="utf-8"))
        self.assertTrue(all(row["address"].startswith("广东省深圳市") for row in rows))
    def test_data_generator_filters_addresses_by_district_and_supports_hong_kong(self):
        mainland = [{"name": "address", "generator": "china_address", "options": {"province": ["广东省"], "city": ["深圳市"], "district": ["南山区"], "street_mode": "none"}}]
        task_id, result = self.runtime.run("data.mock", {"count": 1, "format": "json", "seed": 7, "rules": mainland})
        rows = json.loads((self.temp / "workspace" / task_id / "output" / result.files[0]).read_text(encoding="utf-8"))
        self.assertEqual(rows[0]["address"], "广东省深圳市南山区")
        hong_kong = [{"name": "address", "generator": "china_address", "options": {"province": ["香港特别行政区"], "city": ["香港岛"], "district": ["湾仔区"], "street_mode": "none"}}]
        task_id, result = self.runtime.run("data.mock", {"count": 1, "format": "json", "seed": 7, "rules": hong_kong})
        rows = json.loads((self.temp / "workspace" / task_id / "output" / result.files[0]).read_text(encoding="utf-8"))
        self.assertEqual(rows[0]["address"], "香港特别行政区香港岛湾仔区")
    def test_data_generator_supports_nullable_rate(self):
        rules = [{"name": "optional_text", "generator": "template", "nullable_rate": 1, "options": {"value": "value"}}]
        task_id, result = self.runtime.run("data.mock", {"count": 2, "format": "json", "seed": 7, "rules": rules})
        rows = json.loads((self.temp / "workspace" / task_id / "output" / result.files[0]).read_text(encoding="utf-8"))
        self.assertEqual([row["optional_text"] for row in rows], [None, None])
    def test_data_generator_exports_txt_and_sql(self):
        rules = [{"name": "note", "generator": "template", "options": {"value": "O'Reilly-{index}"}}]
        task_id, result = self.runtime.run("data.mock", {"count": 2, "format": "txt", "seed": 7, "rules": rules, "txt_delimiter": "\t", "txt_header": False})
        self.assertEqual(result.files, [f"{task_id}.txt"])
        txt = (self.temp / "workspace" / task_id / "output" / result.files[0]).read_text(encoding="utf-8")
        self.assertNotIn("record_id\tnote", txt)
        self.assertIn("O'Reilly-1", txt)
        task_id, result = self.runtime.run("data.mock", {"count": 1, "format": "sql", "seed": 7, "rules": rules, "sql_dialect": "postgresql", "sql_table": "public.customer", "sql_batch_size": 1})
        self.assertEqual(result.files, [f"{task_id}.sql"])
        script = (self.temp / "workspace" / task_id / "output" / result.files[0]).read_text(encoding="utf-8")
        self.assertIn('INSERT INTO "public"."customer"', script)
        self.assertIn("O''Reilly-1", script)
        self.assertIn("BEGIN;", script)
    def test_data_generator_exports_zip_bundle(self):
        task_id, result = self.runtime.run("data.mock", {"count": 1, "format": "zip", "seed": 7, "fields": MINIMAL_FIELDS, "zip_formats": ["json", "txt", "sql"]})
        bundle = self.temp / "workspace" / task_id / "output" / result.files[0]
        with __import__("zipfile").ZipFile(bundle) as archive:
            self.assertEqual(set(archive.namelist()), {f"{task_id}.json", f"{task_id}.txt", f"{task_id}.sql", "generation-summary.json"})
        self.assertEqual(result.data["zip_formats"], ["json", "txt", "sql"])
    def test_data_generator_financial_customer_template(self):
        task_id, result = self.runtime.run("data.mock", {"count": 2, "format": "json", "seed": 7, "template": "retail_customer"})
        rows = json.loads((self.temp / "workspace" / task_id / "output" / result.files[0]).read_text(encoding="utf-8"))
        self.assertEqual(rows[0]["customer_id"], "TEST-CUST-00000001")
        self.assertIn("risk_level", rows[0])
        self.assertIn("registered_address", rows[0])
    def test_data_generator_previews_and_selects_multiple_sql_tables(self):
        ddl = self.temp / "multi-table.sql"
        ddl.write_text(
            "CREATE TABLE user_info (id BIGINT PRIMARY KEY, name VARCHAR(30) COMMENT '姓名');\n"
            "CREATE TABLE orders (order_id BIGINT PRIMARY KEY, amount DECIMAL(10,2));",
            encoding="utf-8",
        )
        _, preview = self.runtime.run("data.mock", {"count": 1, "format": "json", "source_file": str(ddl), "source_format": "sql", "preview": True})
        self.assertEqual(preview.status, "success")
        self.assertEqual([table["name"] for table in preview.data["tables"]], ["user_info", "orders"])
        self.assertEqual([field["name"] for field in preview.data["tables"][1]["fields"]], ["order_id", "amount"])
        task_id, result = self.runtime.run("data.mock", {"count": 2, "format": "json", "seed": 7, "source_file": str(ddl), "source_format": "sql", "source_table": "orders"})
        rows = json.loads((self.temp / "workspace" / task_id / "output" / result.files[0]).read_text(encoding="utf-8"))
        self.assertEqual(list(rows[0]), ["order_id", "amount"])
        self.assertEqual(result.data["source_table"], "orders")

    def test_data_generator_rejects_unselected_multiple_sql_tables(self):
        ddl = self.temp / "multi-table.sql"
        ddl.write_text("CREATE TABLE a (id INT); CREATE TABLE b (id INT);", encoding="utf-8")
        _, result = self.runtime.run("data.mock", {"count": 1, "format": "json", "source_file": str(ddl), "source_format": "sql"})
        self.assertEqual(result.status, "failed")
        self.assertIn("source_table", result.message)

    def test_data_generator_groups_excel_structure_by_table(self):
        source = ROOT / "plugins" / "data-generator" / "src" / "main.py"
        spec = importlib.util.spec_from_file_location("data_generator_test_excel_tables", source)
        module = importlib.util.module_from_spec(spec); self.assertIsNotNone(spec.loader); spec.loader.exec_module(module)
        field_list = self.temp / "multi-fields.xlsx"
        field_list.write_bytes(module.xlsx([
            {"table": "users", "field": "id", "type": "INT"},
            {"table": "users", "field": "name", "type": "VARCHAR(20)"},
            {"table": "orders", "field": "order_id", "type": "BIGINT"},
        ]))
        tables = module.excel_tables(field_list)
        self.assertEqual([item["table"] for item in tables], ["users", "orders"])
        self.assertEqual([field["name"] for field in tables[0]["fields"]], ["id", "name"])

    def test_data_generator_infers_rules_from_sql_ddl(self):
        ddl = self.temp / "customer.sql"
        ddl.write_text("CREATE TABLE customer (customer_id VARCHAR(20) COMMENT '客户ID', mobile VARCHAR(11) COMMENT '手机号', amount DECIMAL(18,2) COMMENT '金额');", encoding="utf-8")
        task_id, result = self.runtime.run("data.mock", {"count": 2, "format": "json", "seed": 7, "source_file": str(ddl), "source_format": "sql"})
        rows = json.loads((self.temp / "workspace" / task_id / "output" / result.files[0]).read_text(encoding="utf-8"))
        self.assertEqual(rows[0]["customer_id"], "TEST-CUST-00000001")
        self.assertEqual(len(rows[0]["mobile"]), 11)
        self.assertIsInstance(rows[0]["amount"], float)
    def test_data_generator_infers_rules_from_excel_field_list(self):
        source = ROOT / "plugins" / "data-generator" / "src" / "main.py"
        spec = importlib.util.spec_from_file_location("data_generator_test", source)
        module = importlib.util.module_from_spec(spec); self.assertIsNotNone(spec.loader); spec.loader.exec_module(module)
        field_list = self.temp / "fields.xlsx"
        field_list.write_bytes(module.xlsx([{"field": "mobile", "type": "VARCHAR(11)", "comment": "手机号"}, {"field": "risk_level", "type": "VARCHAR(10)", "comment": "风险等级"}]))
        task_id, result = self.runtime.run("data.mock", {"count": 2, "format": "json", "seed": 7, "source_file": str(field_list), "source_format": "excel"})
        rows = json.loads((self.temp / "workspace" / task_id / "output" / result.files[0]).read_text(encoding="utf-8"))
        self.assertEqual(len(rows[0]["mobile"]), 11)
        self.assertIn(rows[0]["risk_level"], {"低", "中", "高"})
    def test_unknown_parameter_is_rejected(self):
        with self.assertRaises(ValueError): self.runtime.run("data.mock", {"count": 1, "format": "json", "oops": True})
    def test_failed_file_staging_does_not_leave_orphan_workspace(self):
        before = {path.name for path in (self.temp / "workspace").iterdir() if path.is_dir()}
        with self.assertRaises(ValueError):
            self.runtime.run("sql.parse", {"input": str(self.temp / "missing.sql"), "format": "json"})
        after = {path.name for path in (self.temp / "workspace").iterdir() if path.is_dir()}
        self.assertEqual(after, before)

    def test_task_history_and_workspace_clean(self):
        task_id, _ = self.runtime.run("data.mock", {"count": 1, "format": "json", "seed": 7, "fields": MINIMAL_FIELDS})
        record = self.runtime.history.get(task_id)
        self.assertEqual(record["status"], "SUCCEEDED")
        self.assertEqual(record["params"]["seed"], 7)
        self.assertEqual(self.runtime.clean_workspace(date.today() + timedelta(days=1)), 1)
    def test_sensitive_params_are_redacted(self):
        self.assertEqual(Runtime._redact_params({"api_token": "private", "count": 1}), {"api_token": "***", "count": 1})

    def test_sensitive_params_are_redacted_recursively(self):
        value = {"config": {"password": "private"}, "rules": [{"api_key": "secret"}], "count": 1}
        self.assertEqual(Runtime._redact_params(value), {"config": {"password": "***"}, "rules": [{"api_key": "***"}], "count": 1})

    def test_runtime_application_api_exposes_commands_tasks_and_results(self):
        self.assertIn("data.mock", self.runtime.list_commands())
        schema = self.runtime.get_command_schema("data.mock")
        self.assertEqual(schema["type"], "object")
        task_id, result = self.runtime.run("data.mock", {"count": 1, "format": "json", "seed": 7, "fields": MINIMAL_FIELDS})
        self.assertEqual(self.runtime.get_task(task_id)["status"], "SUCCEEDED")
        self.assertEqual(self.runtime.get_task_result(task_id)["status"], "success")
        self.assertEqual(self.runtime.list_tasks(command="data.mock", limit=1)[0]["id"], task_id)
    def test_lock_failure_is_recorded_as_failed_task(self):
        manifest = self.runtime.manager.available["data.mock"]
        manifest.capabilities["concurrency"] = False
        with patch.object(PluginExecutionLock, "acquire", side_effect=OSError("lock unavailable")):
            task_id, result = self.runtime.run("data.mock", {"count": 1, "format": "json", "seed": 7, "fields": MINIMAL_FIELDS})
        self.assertEqual(result.status, "failed")
        self.assertEqual(result.data["error_code"], ErrorCode.CORE_EXECUTION_FAILED)
        self.assertEqual(self.runtime.history.get(task_id)["status"], "FAILED")
        self.assertTrue((self.temp / "workspace" / task_id / "result.json").is_file())

    def test_host_failure_includes_actionable_diagnostics(self):
        source = self.temp / "plugins" / "data-generator" / "src" / "main.py"
        source.write_text(
            "from testbox.sdk import Result\n"
            "class Plugin:\n"
            "    def init(self, context):\n"
            "        raise RuntimeError('intentional host failure')\n"
            "    def execute(self, command, params):\n"
            "        return Result('success', 'unexpected')\n",
            encoding="utf-8",
        )
        _, result = self.runtime.run("data.mock", {"count": 1, "format": "json", "seed": 7, "fields": MINIMAL_FIELDS})
        self.assertEqual(result.status, "failed")
        self.assertEqual(result.data["error_code"], "EXECUTION_FAILED")
        self.assertEqual(result.data["exception_type"], "RuntimeError")
        self.assertIn("intentional host failure", result.data["exception_message"])
        self.assertEqual(result.data["host_exit_code"], 0)
        self.assertIn("RuntimeError: intentional host failure", result.data["task_log_tail"])

    def test_host_timeout_is_recorded(self):
        timed_runtime = Runtime(self.temp, timeout_seconds=0)
        task_id, result = timed_runtime.run("data.mock", {"count": 1, "format": "json", "fields": MINIMAL_FIELDS})
        timed_runtime.close()
        self.assertEqual(result.data["error_code"], "TIMEOUT")
        self.assertEqual(self.runtime.history.get(task_id)["status"], "FAILED")
    def test_incomplete_tasks_are_abandoned_at_startup(self):
        self.runtime.history.create({"id": "20200101T000000-deadbeef", "plugin_name": "data-generator", "plugin_version": "1.0.0", "command": "data.mock", "params": {}, "started_at": "2020-01-01T00:00:00+00:00", "result_path": "missing", "workspace_path": "missing", "host_pid": None})
        recovered = Runtime(self.temp)
        self.assertEqual(recovered.history.get("20200101T000000-deadbeef")["status"], "ABANDONED")
        recovered.close()
    def test_frozen_runtime_uses_user_plugin_directory_for_gui_management(self):
        bundle = self.temp / "bundle"
        shutil.copytree(self.temp / "plugins", bundle / "plugins")
        user_data = self.temp / "user-data"
        archive = self.temp / "sql-select.zip"
        package_plugin(self.temp / "plugins" / "sql-select", archive)

        with patch.object(sys, "frozen", True, create=True), patch.object(sys, "_MEIPASS", str(bundle), create=True), patch.object(Runtime, "_user_data_dir", staticmethod(lambda: user_data)):
            frozen_runtime = Runtime()
            try:
                self.assertEqual(frozen_runtime.plugins_dir, user_data / "plugins")
                self.assertEqual(frozen_runtime.bundled_plugins_dir, bundle / "plugins")
                frozen_runtime.install_plugin(archive)
                self.assertIn("sql.select", frozen_runtime.manager.available)
                frozen_runtime.uninstall_plugin("sql-select")
                self.assertIn("sql.select", frozen_runtime.manager.available)  # 回退到 EXE 内置插件
            finally:
                frozen_runtime.close()

    def test_runtime_plugin_management_refreshes_existing_registry(self):
        archive = self.temp / "data-generator.zip"
        package_plugin(self.temp / "plugins" / "data-generator", archive)
        uninstall_plugin("data-generator", self.temp / "plugins")
        self.runtime.reload_plugins()
        self.assertNotIn("data.mock", self.runtime.manager.available)

        manifest = self.runtime.install_plugin(archive)
        self.assertEqual(manifest.name, "data-generator")
        self.assertIn("data.mock", self.runtime.manager.available)

        self.runtime.uninstall_plugin("data-generator")
        self.assertNotIn("data.mock", self.runtime.manager.available)

    def test_plugin_archive_round_trip(self):
        archive = self.temp / "data-generator.zip"
        package_plugin(self.temp / "plugins" / "data-generator", archive)
        uninstall_plugin("data-generator", self.temp / "plugins")
        after_uninstall = Runtime(self.temp)
        self.assertNotIn("data.mock", after_uninstall.manager.available)
        after_uninstall.close()
        manifest = install_plugin(archive, self.temp / "plugins")
        self.assertEqual(manifest.name, "data-generator")
        after_install = Runtime(self.temp)
        self.assertIn("data.mock", after_install.manager.available)
        after_install.close()

if __name__ == "__main__": unittest.main()
